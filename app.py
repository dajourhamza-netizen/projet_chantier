import io
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from datetime import datetime

# Bibliothèques Word & Excel
from docx import Document
from docxtpl import DocxTemplate, RichText
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
import pandas as pd
import streamlit as st

# ==========================================
# 0. CONFIGURATION DE LA PAGE ET STYLES CSS
# ==========================================
st.set_page_config(
    page_title="Suivi Chantier - Génie Civil & Routes",
    page_icon="🏗️",
    layout="wide",
)

st.markdown(
    """
<style>
    .stApp { background-color: #f8fafc; }
    .gc-header {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        color: #ffffff;
        padding: 22px 28px;
        border-radius: 12px;
        border-left: 8px solid #ff6b00;
        box-shadow: 0px 4px 12px rgba(0, 0, 0, 0.08);
        margin-bottom: 25px;
    }
    .gc-header h1 { color: #ffffff !important; font-size: 26px !important; font-weight: 800 !important; margin: 0 !important; }
    .gc-header p { color: #94a3b8; margin: 6px 0 0 0; font-size: 14px; }
    .kpi-card {
        background-color: #ffffff;
        border: 1px solid #e2e8f0;
        border-top: 4px solid #ff6b00;
        border-radius: 10px;
        padding: 16px;
        text-align: center;
    }
    .kpi-value { font-size: 24px; font-weight: 800; color: #0f172a; }
    .kpi-label { font-size: 11px; font-weight: 600; color: #64748b; text-transform: uppercase; }
    .stButton > button[kind="primary"] {
        background-color: #ff6b00 !important; color: #ffffff !important; border: none !important; border-radius: 8px !important; font-weight: 700 !important;
    }
    section[data-testid="stSidebar"] { background-color: #0f172a !important; color: #ffffff !important; }
    section[data-testid="stSidebar"] label, section[data-testid="stSidebar"] .stMarkdown h1 { color: #f1f5f9 !important; }
</style>
""",
    unsafe_allow_html=True,
)

# ==========================================
# 1. FONCTIONS UTILITAIRES & SYSTEME
# ==========================================

DOSSIER_CHANTIER = (
    os.path.dirname(os.path.abspath(__file__))
    if "__file__" in globals()
    else os.getcwd()
)

chemin_excel_defaut = os.path.join(DOSSIER_CHANTIER, "suivi .xlsx")
if not os.path.exists(chemin_excel_defaut):
  chemin_excel_defaut = os.path.join(DOSSIER_CHANTIER, "suivi.xlsx")

COL_PARTIE = "PARTIE D'OUVRAGE"

COLUMNS_TEMPLATE = [
    "DATE",
    "TITRE DE LA NATURE DES TRAVAUX",
    COL_PARTIE,
    "SITUATION",
    "ACTIVITÉ RÉALISÉE",
    "ÉSSAI/ CONTRÔLE RÉALISÉE",
    "RÉFÉRENCE DE PROCÉDURE",
    "PIÈCES JOINTES",
]

# 📋 Dictionnaire complet
LIAISONS = {
    "ARASE DE PST": {
        "procedure": "TER-PEX-05-00",
        "pieces": (
            "* Fiche de suivi de la PST\n* Fiche de réception topographique\n*"
            " PVs laboratoire"
        ),
    },
    "ARASE DE TERRASSEMENT": {
        "procedure": "TER-PEX-03-00",
        "pieces": (
            "* Fiche de contrôle des déblais\n* Fiche de réception"
            " topographique\n* PVs laboratoire"
        ),
    },
    "ASSISE DE REMBLAIS PURGE": {
        "procedure": "TER-PEX-04-00",
        "pieces": (
            "* Fiche de réception de l'assise des remblais\n* Fiche de"
            " réception topographique\n* Fiche d'identification de la purge\n*"
            " PVs laboratoire"
        ),
    },
    "ASSISE DE REMBLAIS": {
        "procedure": "TER-PEX-04-00",
        "pieces": (
            "* Fiche de réception de l'assise des remblais\n* Fiche de"
            " réception topographique\n* PVs laboratoire"
        ),
    },
    "ASSISE DE REMBLAIS CDF": {
        "procedure": "TER-PEX-04-00",
        "pieces": (
            "* Fiche de réception de l'assise des remblais\n* Fiche de"
            " réception topographique\n* PVs laboratoire"
        ),
    },
    "ASSISE DE REMBLAIS CONTIGUS": {
        "procedure": "OVA-PEX-16-00",
        "pieces": (
            "* Fiche de suivi des remblais contigus\n* Fiche de contrôle des"
            " remblais contigus\n* PVs laboratoire\n* Fiche de réception"
            " topographique"
        ),
    },
    "ASSISE DE REMBLAI DE FOUILLE": {
        "procedure": "OVA-PEX-04-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des fouilles et remblaiement de"
            " fouilles\n* PVs laboratoire"
        ),
    },
    "ASSISE DE REMBLAIS RENFORCE": {
        "procedure": "TER-PEX-13-00",
        "pieces": (
            "* PV Manifold\n* PVs laboratoire\n* Fiche de réception"
            " topographique\n* Fiche de réception assise remblai renforcé"
        ),
    },
    "ASSISE DRAINANTE": {
        "procedure": "TER-PEX-13-00",
        "pieces": (
            "* Fiche de réception topographique\n* PVs laboratoire\n* Fiche de"
            " contrôle de l'assise drainante"
        ),
    },
    "COUCHE DE FORME": {
        "procedure": "TER-PEX-09-00",
        "pieces": (
            "* Fiche de suivi et de contrôle de la CDF\n* Fiche de réception"
            " topographique\n* PVs laboratoire"
        ),
    },
    "DÉCAPAGE": {
        "procedure": "TER-PEX-02-00",
        "pieces": (
            "* Fiche de suivi et de contrôle du décapage\n* Fiche des sections"
            " à décaper\n* Fiche de réception topographique"
        ),
    },
    "DEGAGEMENT D'EMPRISE": {
        "procedure": "TER-PEX-01-00",
        "pieces": (
            "* Fiche de suivi et de contrôle du dégagement des emprises\n*"
            " Fiche de réception topographique\n* Constat dégagement d'emprise"
        ),
    },
    "REMBLAIS": {
        "procedure": "TER-PEX-04-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des remblais\n* PVs laboratoire"
        ),
    },
    "REMBLAIS CDF": {
        "procedure": "TER-PEX-04-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des remblais\n* PVs laboratoire"
        ),
    },
    "REMBLAIS CONTIGUS": {
        "procedure": "OVA-PEX-16-00",
        "pieces": (
            "* Fiche de suivi des remblais contigus\n* Fiche de contrôle des"
            " remblais contigus\n* PVs laboratoire\n* Fiche de réception"
            " topographique"
        ),
    },
    "REMBLAIS DE FOUILLE": {
        "procedure": "OVA-PEX-04-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des fouilles et remblaiement de"
            " fouilles\n* PVs laboratoire"
        ),
    },
    "REMBLAIS DE FOUILLS CDF": {
        "procedure": "OVA-PEX-04-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des fouilles et remblaiement de"
            " fouilles\n* PVs laboratoire"
        ),
    },
    "REMBLAIS RENFORCE": {
        "procedure": "TER-PEX-13-00",
        "pieces": (
            "* Fiche de suivi des remblais renforcé\n* Fiche de contrôle des"
            " armatures Geostrap\n* Fiche de réception de pose des ecailles\n*"
            " PVs laboratoire"
        ),
    },
    "REMBLAIS PST": {
        "procedure": "TER-PEX-05-00",
        "pieces": (
            "* Fiche de suivi et de contrôle des remblais PST\n* PVs"
            " laboratoire"
        ),
    },
}


def text_to_richtext(text):
  if not text or pd.isna(text):
    return ""
  rt = RichText()
  lines = str(text).split("\n")
  for i, line in enumerate(lines):
    rt.add(line)
    if i < len(lines) - 1:
      rt.add("\n")
  return rt


def clean_filename(text):
  if not text:
    return ""
  text = str(text)
  if text.lower().endswith(".docx"):
    text = text[:-5]
  text = (
      unicodedata.normalize("NFD", text)
      .encode("ascii", "ignore")
      .decode("utf-8")
  )
  return re.sub(r"[^a-zA-Z0-9]", "", text).lower()


def trouver_modele_word(nom_nature):
  target_clean = clean_filename(nom_nature)
  if os.path.exists(DOSSIER_CHANTIER):
    for file in os.listdir(DOSSIER_CHANTIER):
      if file.lower().endswith(".docx") and not file.startswith("~$"):
        if clean_filename(file) == target_clean:
          return os.path.join(DOSSIER_CHANTIER, file)
  return None


def construire_nom_pdf(row):
  nature = str(row.get("TITRE DE LA NATURE DES TRAVAUX", "")).strip()
  partie = str(row.get(COL_PARTIE, row.get("PARTIE D meOUVRAGE", ""))).strip()
  situation = str(row.get("SITUATION", "")).strip()
  nom_propre = re.sub(r'[\\/*?:"<>|]', "_", f"{nature} - {partie} - {situation}")
  return f"{re.sub(r'\s+', ' ', nom_propre).strip()}.pdf"


def get_col_val(row, *candidates):
  for c in candidates:
    for col in row.index:
      if str(col).strip().lower() == str(c).strip().lower():
        val = row[col]
        if isinstance(val, (datetime, pd.Timestamp)):
          return val.strftime("%d/%m/%Y")
        val_str = str(val).strip()
        if val_str and val_str.lower() != "nan":
          return val_str
  return ""


def generer_docx_et_pdf_bytes(chemin_modele, contexte):
  with tempfile.TemporaryDirectory() as temp_dir:
    doc = DocxTemplate(chemin_modele)
    doc.render(contexte)
    docx_temp_path = os.path.join(temp_dir, "temp.docx")
    doc.save(docx_temp_path)

    with open(docx_temp_path, "rb") as f:
      docx_bytes = f.read()

    pdf_temp_path = os.path.join(temp_dir, "temp.pdf")
    pdf_bytes = None
    try:
      from docx2pdf import convert

      convert(docx_temp_path, pdf_temp_path)
      if os.path.exists(pdf_temp_path):
        with open(pdf_temp_path, "rb") as f:
          pdf_bytes = f.read()
    except Exception:
      pass

    if pdf_bytes is None:
      pdf_bytes = docx_bytes
    return docx_bytes, pdf_bytes


def generer_di_style_vba(chemin_modele, df_jour):
  doc = Document(chemin_modele)
  word_table = None
  for tbl in doc.tables:
    if len(tbl.columns) >= 4:
      word_table = tbl
      break

  if not word_table:
    raise ValueError("Tableau introuvable dans le Word.")

  for idx, (_, row) in enumerate(df_jour.iterrows()):
    date_val = get_col_val(row, "DATE")
    nature_val = get_col_val(row, "TITRE DE LA NATURE DES TRAVAUX", "NATURE")
    partie_val = get_col_val(
        row, "PARTIE D'OUVRAGE", "PARTIE D meOUVRAGE", "PARTIE"
    )
    situation_val = get_col_val(row, "SITUATION", "PK")
    activite_val = get_col_val(row, "ACTIVITÉ RÉALISÉE", "ACTIVITE")
    essai_val = get_col_val(row, "ÉSSAI/ CONTRÔLE RÉALISÉE", "ESSAI")

    col2_text = (
        f"{activite_val} - {nature_val}"
        if (activite_val and nature_val)
        else (activite_val or nature_val)
    )
    col3_text = (
        f"{partie_val} / {situation_val}"
        if (partie_val and situation_val)
        else (partie_val or situation_val)
    )

    target_row_idx = idx + 1
    if target_row_idx < len(word_table.rows):
      row_cells = word_table.rows[target_row_idx].cells
    else:
      row_cells = word_table.add_row().cells

    row_cells[0].text = date_val
    row_cells[1].text = col2_text
    row_cells[2].text = col3_text
    row_cells[3].text = essai_val
    if len(row_cells) >= 5:
      row_cells[4].text = ""

  return doc


def get_sheet_names(filepath):
  if os.path.exists(filepath):
    try:
      return pd.ExcelFile(filepath).sheet_names
    except Exception:
      return ["Chantier Principal"]
  return ["Chantier Principal"]


def save_to_excel_with_formatting(
    df_to_save, filepath, sheet_name="Chantier Principal"
):
  try:
    df_clean = df_to_save.copy()
    if "Imprimer" in df_clean.columns:
      df_clean = df_clean.drop(columns=["Imprimer"])

    if "DATE" in df_clean.columns:
      df_clean["DATE"] = pd.to_datetime(
          df_clean["DATE"], dayfirst=True, errors="coerce"
      )
      df_clean = df_clean.sort_values(by="DATE", ascending=True)

    mode = "a" if os.path.exists(filepath) else "w"
    kwargs = {"mode": mode, "engine": "openpyxl"}
    if mode == "a":
      kwargs["if_sheet_exists"] = "replace"

    with pd.ExcelWriter(filepath, **kwargs) as writer:
      df_clean.to_excel(writer, sheet_name=sheet_name, index=False)
      worksheet = writer.sheets[sheet_name]

      for col_idx, column in enumerate(worksheet.columns, start=1):
        max_len = max(len(str(cell.value or "")) for cell in column)
        col_letter = column[0].column_letter
        worksheet.column_dimensions[col_letter].width = min(max_len + 4, 60)

        if column[0].value == "DATE":
          for cell in column[1:]:
            if cell.value:
              cell.number_format = "DD/MM/YYYY"

      num_rows = max(len(df_clean) + 1, 2)
      num_cols = len(df_clean.columns)
      if num_cols > 0:
        end_col = get_column_letter(num_cols)
        clean_sheet = re.sub(r"\W+", "_", sheet_name)
        worksheet._tables.clear()
        tab = Table(
            displayName=f"Tab_{clean_sheet}", ref=f"A1:{end_col}{num_rows}"
        )
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium3", showRowStripes=True
        )
        worksheet.add_table(tab)

    return True, "✅ Fichier Excel mis à jour et trié par date avec succès !"
  except PermissionError:
    return (
        False,
        "❌ **Erreur d'accès !** Le fichier Excel est ouvert sur votre"
        " ordinateur. Veuillez le FERMER puis réessayez.",
    )
  except Exception as e:
    return False, f"❌ Erreur : {e}"


# ==========================================
# 2. BARRE LATÉRALE (SIDEBAR) & GESTION PROJETS
# ==========================================
st.sidebar.markdown("### 🏗️ **Gestion de Chantier**")

chantiers_existants = get_sheet_names(chemin_excel_defaut)
chantier_actif = st.sidebar.selectbox(
    "📌 **Projet Actif :**", options=chantiers_existants
)

st.sidebar.markdown("---")

# ➕ Section Création d'un Nouveau Projet
with st.sidebar.expander(
    "➕ **Créer / Ajouter un Nouveau Projet**", expanded=False
):
  nouveau_projet_nom = st.text_input(
      "Nom du nouveau projet / chantier :", key="new_proj_input"
  )
  if st.button(
      "✨ Créer le Projet",
      type="primary",
      key="btn_create_proj",
      use_container_width=True,
  ):
    nom_clean = nouveau_projet_nom.strip()
    if nom_clean:
      if nom_clean in chantiers_existants:
        st.sidebar.error("⚠️ Ce projet existe déjà !")
      else:
        df_vide = pd.DataFrame(columns=COLUMNS_TEMPLATE)
        success, msg = save_to_excel_with_formatting(
            df_vide, chemin_excel_defaut, sheet_name=nom_clean
        )
        if success:
          st.sidebar.success(f"✅ Projet '{nom_clean}' créé !")
          st.rerun()
        else:
          st.sidebar.error(msg)
    else:
      st.sidebar.warning("⚠️ Veuillez entrer un nom valide.")

# Chargement des données du projet actif
df = None
if os.path.exists(chemin_excel_defaut):
  try:
    df = pd.read_excel(chemin_excel_defaut, sheet_name=chantier_actif).fillna(
        ""
    )
    if "DATE" in df.columns:
      df["DATE"] = pd.to_datetime(df["DATE"], dayfirst=True, errors="coerce")
      df = df.sort_values(by="DATE", ascending=True)
      df["DATE"] = df["DATE"].dt.strftime("%d/%m/%Y").fillna("")
    st.sidebar.success(f"✅ Projet '{chantier_actif}' chargé ({len(df)} lignes)")
  except Exception as e:
    st.sidebar.error(f"❌ Erreur Excel : {e}")

# ==========================================
# 3. INTERFACE PRINCIPALE
# ==========================================

st.markdown(
    f"""
<div class="gc-header">
    <h1>🛣️ Plateforme Génie Civil & Travaux Routiers</h1>
    <p>Projet Actif : <b>{chantier_actif}</b></p>
</div>
""",
    unsafe_allow_html=True,
)

if df is not None:
  tab1, tab2, tab3 = st.tabs([
      "📝 **Nouvelle Saisie Chantier**",
      "📊 **Registre & Génération Individuelle**",
      "📅 **Demandes d'Intervention (DI) Multi-Dates**",
  ])

  # -------------------------------------------------------------
  # TAB 1 : SAISIE (MODIFIÉ AVEC LISTE DYNAMIQUE PARTIE D'OUVRAGE)
  # -------------------------------------------------------------
  with tab1:
    st.markdown("##### 👷 **Ajouter une nouvelle fiche de contrôle**")
    col1, col2 = st.columns(2)
    with col1:
      date_saisie = st.date_input(
          "🗓️ Date des Travaux", value=datetime.today(), format="DD/MM/YYYY"
      )
      nature_selectionnee = st.selectbox(
          "📌 Nature des travaux", options=list(LIAISONS.keys())
      )
      info_liaison = LIAISONS.get(
          nature_selectionnee, {"procedure": "", "pieces": ""}
      )

      # 🧱 DÉBUT DE LA MODIFICATION DE PARTIE D'OUVRAGE DYNAMIQUE
      parties_existantes = []
      if df is not None and COL_PARTIE in df.columns:
        parties_existantes = sorted(
            list(
                set(
                    [
                        str(p).strip()
                        for p in df[COL_PARTIE].unique()
                        if str(p).strip() and str(p).lower() != "nan"
                    ]
                )
            )
        )

      options_partie = parties_existantes + ["➕ Autre / Nouvelle partie..."]

      partie_choisie = st.selectbox(
          "🧱 Partie d'ouvrage (Sélectionner ou Ajouter)",
          options=options_partie,
      )

      if partie_choisie == "➕ Autre / Nouvelle partie...":
        partie_ouvrage = st.text_input(
            "✍️ Saisir la nouvelle Partie d'ouvrage :",
            placeholder="Ex: CULEE C0...",
        )
      else:
        partie_ouvrage = partie_choisie
      # 🧱 FIN DE LA MODIFICATION

      situation = st.text_input(
          "📍 Situation / PK", placeholder="Ex: PK 1+120 AU PK 1+220"
      )

    with col2:
      activite = st.text_area("🚜 Activité réalisée", height=80)

      essai = st.selectbox(
          "🧪 Essai / Contrôle réalisé",
          options=[
              "Aucun",
              "TENEUR EN EAU",
              "CAMPACITÉ",
              "ESSAI À LA PLAQUE",
              "ESSAI À LA PLAQUE + CAMPACITÉ",
              "PRELEVEMENT APRES COMPACTAGE",
              "PRELEVEMENT AVANT COMPACTAGE",
              "IDENTIFICATION DES MATERIAUX",
              "PRELEVEMENT",
          ],
      )

      procedure = st.text_input(
          "📑 Référence procédure", value=info_liaison["procedure"]
      )
      pieces_jointes = st.text_area(
          "📎 Pièces jointes", value=info_liaison["pieces"], height=100
      )

    if st.button("💾 Enregistrer la Fiche", type="primary"):
      new_entry = {
          "DATE": date_saisie.strftime("%d/%m/%Y"),
          "TITRE DE LA NATURE DES TRAVAUX": nature_selectionnee,
          COL_PARTIE: partie_ouvrage,
          "SITUATION": situation,
          "ACTIVITÉ RÉALISÉE": activite,
          "ÉSSAI/ CONTRÔLE RÉALISÉE": "" if essai == "Aucun" else essai,
          "RÉFÉRENCE DE PROCÉDURE": procedure,
          "PIÈCES JOINTES": pieces_jointes,
      }
      df_updated = pd.concat(
          [df, pd.DataFrame([new_entry])], ignore_index=True
      )

      success, msg = save_to_excel_with_formatting(
          df_updated, chemin_excel_defaut, sheet_name=chantier_actif
      )
      if success:
        st.success(msg)
        st.rerun()
      else:
        st.error(msg)

  # -------------------------------------------------------------
  # TAB 2 : REGISTRE & FILTRES COMPLETS
  # -------------------------------------------------------------
  with tab2:
    st.markdown("##### 🔍 **Registre des Travaux, Filtrage & Exportation**")

    # 🎯 SECTION FILTRES DE RECHERCHE
    with st.expander("🎯 **Filtres de Recherche & Tri**", expanded=True):
      col_f1, col_f2, col_f3, col_f4 = st.columns(4)

      with col_f1:
        natures_dispo = ["Toutes"] + list(
            sorted([
                str(n)
                for n in df["TITRE DE LA NATURE DES TRAVAUX"].unique()
                if str(n).strip() and str(n).lower() != "nan"
            ])
        )
        filtre_nature = st.selectbox("📌 Nature :", options=natures_dispo)

      with col_f2:
        dates_dispo = ["Toutes"] + sorted([
            str(d).strip()
            for d in df["DATE"].unique()
            if str(d).strip() and str(d).lower() != "nan"
        ])
        filtre_date = st.selectbox("🗓️ Date :", options=dates_dispo)

      with col_f3:
        parties_dispo = ["Toutes"] + list(
            sorted([
                str(p).strip()
                for p in df[COL_PARTIE].unique()
                if str(p).strip() and str(p).lower() != "nan"
            ])
        )
        filtre_partie = st.selectbox(
            "🧱 Partie d'ouvrage :", options=parties_dispo
        )

      with col_f4:
        recherche_texte = st.text_input(
            "🔍 Recherche globale :", placeholder="PK, Activité..."
        )

    # Application des filtres
    df_filtre = df.copy()

    if filtre_nature != "Toutes":
      df_filtre = df_filtre[
          df_filtre["TITRE DE LA NATURE DES TRAVAUX"].astype(str).str.strip()
          == filtre_nature
      ]

    if filtre_date != "Toutes":
      df_filtre = df_filtre[
          df_filtre["DATE"].astype(str).str.strip() == filtre_date
      ]

    if filtre_partie != "Toutes":
      df_filtre = df_filtre[
          df_filtre[COL_PARTIE].astype(str).str.strip() == filtre_partie
      ]

    if recherche_texte.strip():
      mots = recherche_texte.strip().lower()
      mask = (
          df_filtre[COL_PARTIE].astype(str).str.lower().str.contains(mots)
          | df_filtre["SITUATION"].astype(str).str.lower().str.contains(mots)
          | df_filtre["ACTIVITÉ RÉALISÉE"]
          .astype(str)
          .str.lower()
          .str.contains(mots)
      )
      df_filtre = df_filtre[mask]

    st.caption(
        f"📊 Résultats affichés : **{len(df_filtre)}** / {len(df)} fiches"
    )

    df_editor = df_filtre.copy()
    if "Imprimer" not in df_editor.columns:
      df_editor.insert(0, "Imprimer", False)

    edited_df = st.data_editor(
        df_editor, num_rows="dynamic", height=400, use_container_width=True
    )

    st.markdown("---")
    col_act1, col_act2 = st.columns(2)

    with col_act1:
      if st.button(
          "💾 Enregistrer les modifications dans Excel",
          type="secondary",
          use_container_width=True,
      ):
        df_sauvegarde = df.copy()
        edited_clean = edited_df.drop(columns=["Imprimer"], errors="ignore")
        df_sauvegarde.loc[edited_clean.index] = edited_clean

        success, msg = save_to_excel_with_formatting(
            df_sauvegarde, chemin_excel_defaut, sheet_name=chantier_actif
        )
        if success:
          st.success(msg)
          st.rerun()
        else:
          st.error(msg)

    with col_act2:
      lignes_selectionnees = edited_df[edited_df["Imprimer"] == True]
      nb_selections = len(lignes_selectionnees)

      if st.button(
          f"📦 Générer les Fiches Sélectionnées ({nb_selections})",
          type="primary",
          use_container_width=True,
      ):
        if nb_selections == 0:
          st.warning("⚠️ Cochez au moins une case 'Imprimer' dans le tableau.")
        else:
          zip_buffer = io.BytesIO()
          fichiers_crees = 0
          with zipfile.ZipFile(
              zip_buffer, "w", zipfile.ZIP_DEFLATED
          ) as zip_file:
            for idx, row in lignes_selectionnees.iterrows():
              nom_modele = get_col_val(
                  row, "TITRE DE LA NATURE DES TRAVAUX", "NATURE"
              )
              chemin_modele = trouver_modele_word(nom_modele)
              if chemin_modele:
                contexte = {
                    "NATURE": get_col_val(
                        row, "TITRE DE LA NATURE DES TRAVAUX", "NATURE"
                    ),
                    "REF": get_col_val(row, "RÉFÉRENCE DE PROCÉDURE", "REF"),
                    "PARTIE": get_col_val(
                        row,
                        "PARTIE D'OUVRAGE",
                        "PARTIE D meOUVRAGE",
                        "PARTIE",
                    ),
                    "SITUATION": get_col_val(row, "SITUATION", "PK"),
                    "PIECES": text_to_richtext(
                        get_col_val(row, "PIÈCES JOINTES", "PIECES")
                    ),
                    "DATE": get_col_val(row, "DATE"),
                    "ACTIVITE": text_to_richtext(
                        get_col_val(row, "ACTIVITÉ RÉALISÉE", "ACTIVITE")
                    ),
                    "ESSAI": get_col_val(
                        row, "ÉSSAI/ CONTRÔLE RÉALISÉE", "ESSAI"
                    ),
                }
                docx_b, pdf_b = generer_docx_et_pdf_bytes(
                    chemin_modele, contexte
                )
                nom_base = construire_nom_pdf(row).replace(".pdf", "")
                zip_file.writestr(f"{nom_base}.docx", docx_b)
                zip_file.writestr(f"{nom_base}.pdf", pdf_b)
                fichiers_crees += 1

          if fichiers_crees > 0:
            zip_buffer.seek(0)
            st.download_button(
                label="📦 Télécharger Pack ZIP",
                data=zip_buffer,
                file_name="Fiches_Chantier.zip",
                mime="application/zip",
                use_container_width=True,
            )

  # -------------------------------------------------------------
  # TAB 3 : DI MULTI-DATES (EXPORT PDF)
  # -------------------------------------------------------------
  with tab3:
    st.markdown(
        "##### 📅 **Génération des Demandes d'Intervention (DI) en PDF**"
    )
    dates_disponibles = sorted([
        str(d).strip()
        for d in df["DATE"].unique()
        if str(d).strip() and str(d).lower() != "nan"
    ])

    dates_choisies = st.multiselect(
        "🗓️ Sélectionner une ou plusieurs dates :", options=dates_disponibles
    )

    if dates_choisies and st.button(
        "📑 Générer DI Globale en PDF", type="primary"
    ):
      modele_di = os.path.join(DOSSIER_CHANTIER, "Demande d'intervention.docx")
      if not os.path.exists(modele_di):
        modele_di = os.path.join(DOSSIER_CHANTIER, "Demande_intervention.docx")

      if os.path.exists(modele_di):
        zip_buffer = io.BytesIO()
        has_pdf = False

        with zipfile.ZipFile(
            zip_buffer, "w", zipfile.ZIP_DEFLATED
        ) as zip_file:
          for d_single in dates_choisies:
            df_sub = df[df["DATE"].astype(str).str.strip() == d_single]
            if not df_sub.empty:
              doc_rempli = generer_di_style_vba(modele_di, df_sub)

              with tempfile.TemporaryDirectory() as temp_dir:
                docx_path = os.path.join(temp_dir, "temp_di.docx")
                pdf_path = os.path.join(temp_dir, "temp_di.pdf")

                doc_rempli.save(docx_path)
                pdf_bytes = None

                # 1. Essai avec docx2pdf (Windows / Word)
                try:
                  from docx2pdf import convert

                  convert(docx_path, pdf_path)
                  if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as pf:
                      pdf_bytes = pf.read()
                except Exception:
                  pass

                # 2. Essai avec LibreOffice (Linux / Server)
                if pdf_bytes is None:
                  try:
                    subprocess.run(
                        [
                            "soffice",
                            "--headless",
                            "--convert-to",
                            "pdf",
                            "--outdir",
                            temp_dir,
                            docx_path,
                        ],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        check=True,
                    )
                    if os.path.exists(pdf_path):
                      with open(pdf_path, "rb") as pf:
                        pdf_bytes = pf.read()
                  except Exception:
                    pass

                nom_fichier = f"DI_Globale_{d_single.replace('/', '-')}"

                if pdf_bytes:
                  zip_file.writestr(f"{nom_fichier}.pdf", pdf_bytes)
                  has_pdf = True
                else:
                  with open(docx_path, "rb") as f_docx:
                    zip_file.writestr(f"{nom_fichier}.docx", f_docx.read())

        zip_buffer.seek(0)
        st.download_button(
            label="📦 Télécharger ZIP des DI (PDF)",
            data=zip_buffer,
            file_name="DI_Globales_PDF.zip",
            mime="application/zip",
            use_container_width=True,
        )

        if not has_pdf:
          st.warning(
              "⚠️ Microsoft Word ou LibreOffice n'a pas pu être exécuté sur la"
              " machine pour convertir en PDF. Le fichier a été généré en Word"
              " (.docx) comme secours."
          )
      else:
        st.error(
            "❌ Modèle `Demande d'intervention.docx` introuvable dans le dossier"
            " du projet."
        )
